from __future__ import annotations

import re
import uuid
from pathlib import Path
from typing import Any

from docx import Document
from docx.document import Document as DocxDocument
from docx.table import _Cell, Table
from docx.text.paragraph import Paragraph
from docx.text.run import Run

from .documents import STORAGE_ROOT


class RenderableValue:
    def __init__(self, value: Any, confidence: float = 1.0, needs_review: bool = False):
        self.value = value
        self.confidence = confidence
        self.needs_review = needs_review

    def __str__(self) -> str:
        return str(self.value)


class ReportGenerator:
    def generate_docx(self, template_path: str | None, template_content_json: dict[str, Any], field_values: dict[str, Any], output_path: Path, header_image_path: Path | None = None, certificate_text: str | None = None) -> Path:
        resolved_template_path = self._resolve_template_path(template_path)
        if resolved_template_path is None:
            raise ValueError("original_docx_url is required to generate a template-preserving report")

        # Load template for rendering and clone reference for structural benchmark validation
        template_doc = Document(str(resolved_template_path))
        document = Document(str(resolved_template_path))

        from .docx import HeaderEngine, CertificateEngine, TemplateRenderer, ValidationEngine

        # Phase 2: Header Graphic Engine
        HeaderEngine.replace_header(document, header_image_path)

        # Collect known template tokens from DB template content JSON
        known_tokens: list[str] = []
        for field in self._iter_template_fields(template_content_json):
            fc = self._clean_value(field.get("field_code"))
            lbl = self._clean_value(field.get("label"))
            dyn = self._clean_value(field.get("dynamic_value"))
            if fc: known_tokens.append(fc)
            if lbl:
                known_tokens.append(lbl)
                slug = self._slugify(lbl)
                if slug: known_tokens.append(slug)
            if dyn: known_tokens.append(dyn)

        reserved_headers = {"HEADER", "HEADER_IMAGE", "header", "header_image", "HEADER_LOGO", "header_logo"}
        for key in field_values.keys():
            cleaned_key = self._clean_value(key)
            if cleaned_key and cleaned_key not in reserved_headers:
                known_tokens.append(cleaned_key)

        known_tokens = [t for t in known_tokens if t not in reserved_headers]

        # Render Completion Certificate (before valuation tables)
        CertificateEngine.render_certificate(document, certificate_text, known_tokens, field_values)

        # Phase 1: In-Place Character-Stream Template Renderer
        TemplateRenderer.render_document(document, known_tokens, field_values)

        # Phase 3: Validation Engine & Structural Guards
        ValidationEngine.validate_generation(template_doc, document, known_tokens)

        final_output_path = self._resolve_output_path(output_path)
        final_output_path.parent.mkdir(parents=True, exist_ok=True)
        document.save(str(final_output_path))
        return final_output_path

    def _resolve_template_path(self, template_path: str | None) -> Path | None:
        if not template_path:
            return None
        if template_path.startswith("/storage/"):
            return STORAGE_ROOT / template_path.removeprefix("/storage/")
        return Path(template_path)

    def _resolve_output_path(self, output_path: Path) -> Path:
        if output_path.parent.name == "generated_reports":
            return output_path

        report_id = output_path.stem or uuid.uuid4().hex
        return STORAGE_ROOT / "generated_reports" / f"{report_id}.docx"

    def _resolve_field_values(self, template_content_json: dict[str, Any], field_values: dict[str, Any]) -> dict[str, str]:
        resolved: dict[str, str] = {}

        for field in self._iter_template_fields(template_content_json):
            field_code = self._clean_value(field.get("field_code"))
            label = self._clean_value(field.get("label"))
            dynamic_value_key = self._clean_value(field.get("dynamic_value"))
            static_value = self._clean_value(field.get("static_value"))

            candidates = [field_code, label, self._slugify(label), dynamic_value_key]
            chosen_value = None
            for candidate in candidates:
                if candidate and candidate in field_values and field_values[candidate] not in (None, ""):
                    chosen_value = field_values[candidate]
                    break

            if chosen_value in (None, ""):
                chosen_value = static_value

            if chosen_value not in (None, ""):
                if field_code:
                    resolved[field_code] = str(chosen_value)
                if label:
                    resolved[label] = str(chosen_value)
                slug_label = self._slugify(label)
                if slug_label:
                    resolved[slug_label] = str(chosen_value)

                for keyword in self._coerce_keywords(field.get("keywords", [])):
                    resolved[keyword] = str(chosen_value)

        for key, value in field_values.items():
            if value not in (None, ""):
                resolved[self._clean_value(key)] = str(value)

        return resolved

    def _iter_template_fields(self, template_content_json: dict[str, Any]) -> list[dict[str, Any]]:
        fields: list[dict[str, Any]] = []

        def walk_section(section: dict[str, Any]) -> None:
            for field in section.get("fields", []) or []:
                fields.append(field)
                nested_fields = field.get("nested_fields") or []
                if nested_fields:
                    walk_nested(nested_fields)

        def walk_nested(nodes: list[dict[str, Any]]) -> None:
            for node in nodes:
                fields.append(node)
                nested_fields = node.get("nested_fields") or []
                if nested_fields:
                    walk_nested(nested_fields)

        for section in template_content_json.get("sections", []) or []:
            walk_section(section)
        return fields

    def _replace_placeholders(self, document: DocxDocument, field_values: dict[str, str]) -> None:
        for paragraph in self._iter_paragraphs(document):
            self._replace_in_paragraph(paragraph, field_values)

        for table in document.tables:
            self._replace_in_table(table, field_values)

        for section in document.sections:
            for paragraph in self._iter_paragraphs(section.header):
                self._replace_in_paragraph(paragraph, field_values)
            for table in section.header.tables:
                self._replace_in_table(table, field_values)
            for paragraph in self._iter_paragraphs(section.footer):
                self._replace_in_paragraph(paragraph, field_values)
            for table in section.footer.tables:
                self._replace_in_table(table, field_values)

    def _iter_paragraphs(self, container: Any):
        if hasattr(container, "paragraphs"):
            for paragraph in container.paragraphs:
                yield paragraph
        if hasattr(container, "tables"):
            for table in container.tables:
                yield from self._iter_table_paragraphs(table)

    def _iter_table_paragraphs(self, table: Table):
        for row in table.rows:
            for cell in row.cells:
                for paragraph in cell.paragraphs:
                    yield paragraph
                for nested_table in cell.tables:
                    yield from self._iter_table_paragraphs(nested_table)

    def _replace_in_table(self, table: Table, field_values: dict[str, str]) -> None:
        for row in table.rows:
            for cell in row.cells:
                self._replace_in_cell(cell, field_values)

    def _replace_in_cell(self, cell: _Cell, field_values: dict[str, str]) -> None:
        for paragraph in cell.paragraphs:
            self._replace_in_paragraph(paragraph, field_values)
        for table in cell.tables:
            self._replace_in_table(table, field_values)

    def _replace_in_paragraph(self, paragraph: Paragraph, field_values: dict[str, str]) -> None:
        for run in paragraph.runs:
            self._replace_in_run(run, field_values)

    def _replace_in_run(self, run: Run, field_values: dict[str, str]) -> None:
        if not run.text:
            return

        original_text = run.text
        updated_text = original_text
        for token, value in field_values.items():
            updated_text = self._replace_token_variants(updated_text, token, value)

        if updated_text != original_text:
            run.text = updated_text

    def _replace_token_variants(self, text: str, token: str, value: str) -> str:
        replacements = [
            f"{{{{{token}}}}}",
            f"{{{token}}}",
            f"[{token}]",
            f"<{token}>",
            token,
        ]
        updated_text = text
        for pattern in replacements:
            updated_text = updated_text.replace(pattern, value)
        return updated_text

    def _clean_value(self, value: Any) -> str:
        return re.sub(r"\s+", " ", str(value or "")).strip()

    def _slugify(self, value: Any) -> str:
        cleaned = self._clean_value(value).lower()
        return re.sub(r"[^a-z0-9]+", "_", cleaned).strip("_")

    def _coerce_keywords(self, keywords: Any) -> list[str]:
        if not isinstance(keywords, list):
            return []
        return [self._clean_value(keyword) for keyword in keywords if self._clean_value(keyword)]