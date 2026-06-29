from __future__ import annotations

from pathlib import Path
from docx.document import Document as DocxDocument
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Inches


class HeaderRenderer:
    @staticmethod
    def replace_header(document: DocxDocument, header_image_path: Path | None) -> None:
        if not header_image_path or not header_image_path.exists() or not header_image_path.is_file():
            return

        image_str = str(header_image_path)
        printable_width = Inches(6.5)

        # Strategy 1: Word Section Header (Section.header)
        section = document.sections[0] if document.sections else None
        if section and section.header:
            header = section.header
            # Check if header has drawings or text
            has_header_content = False
            for p in header.paragraphs:
                if p.text.strip() or "drawing" in p._element.xml.lower():
                    has_header_content = True
                    break
            if not has_header_content:
                for t in header.tables:
                    if t._element.xml:
                        has_header_content = True
                        break

            if has_header_content:
                # Replace content in Section.header
                for p in list(header.paragraphs):
                    p.text = ""
                p = header.paragraphs[0] if header.paragraphs else header.add_paragraph()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(image_str, width=printable_width)
                return

        # Strategy 2: First-Page Body Branding Block (inspect top tables or paragraphs with drawings)
        # Check first paragraph or first table in document body
        for p in document.paragraphs[:3]:
            if "drawing" in p._element.xml.lower():
                p.text = ""
                # remove existing xml elements in run if needed
                p._element.clear_content()
                p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                r = p.add_run()
                r.add_picture(image_str, width=printable_width)
                return

        if document.tables:
            first_table = document.tables[0]
            # Check if top table is a branding header box
            if len(first_table.rows) <= 3:
                # Check for drawing XML in first row
                for cell in first_table.rows[0].cells:
                    if "drawing" in cell._element.xml.lower():
                        cell.text = ""
                        p = cell.paragraphs[0] if cell.paragraphs else cell.add_paragraph()
                        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
                        r = p.add_run()
                        r.add_picture(image_str, width=printable_width)
                        return

        # Strategy 3: Fallback Centered Insertion at Top of Body
        if document.paragraphs:
            p = document.paragraphs[0].insert_paragraph_before()
        else:
            p = document.add_paragraph()

        p.alignment = WD_ALIGN_PARAGRAPH.CENTER
        r = p.add_run()
        r.add_picture(image_str, width=printable_width)
